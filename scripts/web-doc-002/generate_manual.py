from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
MANIFEST = ROOT / "projects" / "test-workflow" / "artifacts" / "web-doc" / "system-manual" / "manifest.json"
DEFAULT_OUTPUT = ROOT / "projects" / "test-workflow" / "artifacts" / "web-doc-002" / "manual-draft.docx"

# compact_reference_guide tokens, with the user-required A4 page size as a named override.
NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(92, 100, 110)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F7FA"
PAGE_WIDTH_IN = 8.27
PAGE_HEIGHT_IN = 11.69
MARGIN_IN = 1.0
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - 2 * MARGIN_IN


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in):
    widths = [int(round(width * 1440)) for width in widths_in]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D7DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text, *, style=None, size=11, color=None, bold=False, italic=False, align=None, before=0, after=6, line=1.25):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    set_paragraph_spacing(paragraph, before, after, line)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4, line=1.25)
    label_run = p.add_run(f"{label}：")
    set_run_font(label_run, bold=True, color=NAVY)
    value_run = p.add_run(value)
    set_run_font(value_run)
    return p


def add_caption(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=2, after=10, line=1.0)
    run = p.add_run(f"图 {number}  {text}")
    set_run_font(run, size=9, color=GRAY, italic=True)
    return p


def add_image(doc, image_path, number, caption, content_root):
    absolute = content_root / image_path
    if not absolute.exists():
        raise FileNotFoundError(absolute)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=3, after=0, line=1.0)
    run = p.add_run()
    run.add_picture(str(absolute), width=Inches(CONTENT_WIDTH_IN - 0.12))
    add_caption(doc, number, caption)


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_IN])
    set_table_borders(table, color="B8CBE0", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.2)
    r = p.add_run(f"{title}：")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=3, line=1.0)


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 14, 7), ("Heading 3", 12, NAVY, 10, 5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=2, line=1.0)
    run = header.add_run("RSS调度系统操作手册")
    set_run_font(run, size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, before=2, line=1.0)
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    add_page_number(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=GRAY)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(p, before={1: 18, 2: 14, 3: 10}[level], after={1: 10, 2: 7, 3: 5}[level], line=1.25)
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: NAVY}[level], bold=True)
    return p


def build_document(payload, output: Path):
    doc = Document()
    configure_styles(doc)
    content_root = MANIFEST.parent
    used_images: set[str] = set()
    figure_no = 0

    # Cover page.
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(115)
    cover.paragraph_format.space_after = Pt(16)
    run = cover.add_run(payload["documentTitle"])
    set_run_font(run, size=30, color=NAVY, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub, after=18, line=1.0)
    run = sub.add_run("登录与系统管理")
    set_run_font(run, size=16, color=BLUE, bold=True)
    for label, value in (("版本", payload["documentVersion"]), ("生成日期", payload["generatedDate"]), ("适用范围", "系统登录、系统管理")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=4, line=1.0)
        r = p.add_run(f"{label}：{value}")
        set_run_font(r, size=11, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "文档信息", 1)
    table = doc.add_table(rows=5, cols=2)
    set_table_geometry(table, [1.65, CONTENT_WIDTH_IN - 1.65])
    set_table_borders(table)
    rows = [("文档名称", payload["documentTitle"]), ("适用系统", "RSS调度系统"), ("适用范围", "登录、系统管理"), ("文档版本", payload["documentVersion"]), ("环境说明", "基于当前真实测试环境生成")]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, after=0, line=1.15)
        r = p.add_run(label)
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.15)
        r = p.add_run(value)
        set_run_font(r, size=10.5)

    add_heading(doc, "目录", 1)
    toc = doc.add_paragraph("[[TOC]]")
    set_paragraph_spacing(toc, after=12, line=1.15)
    toc.runs[0].font.color.rgb = GRAY
    doc.add_page_break()

    add_heading(doc, "1. 文档说明", 1)
    add_text(doc, "本手册用于说明当前真实 Web 系统中已经实际验证的登录和系统管理操作。正文只保留已确认的页面入口、按钮名称、必要输入和实际结果；未形成合格执行证据的功能不纳入正式步骤。")
    add_note(doc, "账号说明", "请使用系统管理员分配的用户名和密码登录。手册不记录具体账号密码。")
    add_note(doc, "操作提示", "新增、修改、删除或权限调整前，请先确认当前页面、目标对象和必要字段；删除类操作需要在确认提示中再次确认。")

    add_heading(doc, "2. 系统登录", 1)
    login_steps = [s for s in payload["steps"] if s["chapter"] == "系统登录"]
    for step in login_steps:
        add_heading(doc, step["stepId"] + "  " + step["title"], 2)
        add_label_value(doc, "操作", step["action"])
        add_label_value(doc, "预期结果", step["expectedResult"])
        image = step["annotatedScreenshot"]
        if image and image not in used_images:
            figure_no += 1
            add_image(doc, image, figure_no, step["title"], content_root)
            used_images.add(image)

    add_heading(doc, "3. 系统管理", 1)
    add_text(doc, "登录后，在导航中展开“系统管理”。当前真实页面确认的子菜单包括用户管理、角色管理、菜单管理、字典管理和外部系统配置。")
    system_steps = [s for s in payload["steps"] if s["chapter"] == "系统管理"]
    current_feature = None
    feature_count = 0
    for step in system_steps:
        if step["feature"] != current_feature:
            current_feature = step["feature"]
            feature_count += 1
            add_heading(doc, f"3.{feature_count} {current_feature}", 2)
        add_heading(doc, step["stepId"] + "  " + step["title"], 3)
        add_label_value(doc, "操作", step["action"])
        add_label_value(doc, "预期结果", step["expectedResult"])
        image = step["annotatedScreenshot"]
        if image and image not in used_images:
            figure_no += 1
            add_image(doc, image, figure_no, f"{current_feature} - {step['title']}", content_root)
            used_images.add(image)

    add_heading(doc, "4. 常见操作说明与范围", 1)
    add_text(doc, "本手册中的页面截图来自当前真实运行页面。截图中的非敏感示例数据仅用于帮助识别页面位置，不代表业务上必须使用同名数据。")
    add_text(doc, "字典管理的删除操作未列入本手册：已有真实页面执行结果显示当前页面拒绝删除系统类型。部门管理、参数配置、日志管理、登录失败、退出、密码重置、导入导出等也未列入，因为本阶段没有形成相应的合格真实操作证据。")
    add_note(doc, "后续扩展", "如需补充上述功能，应先完成对应真实页面操作和结果确认，再单独增加文档截图与步骤。")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(json.dumps({"output": str(output), "figures": figure_no, "uniqueImages": len(used_images)}, ensure_ascii=False))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", default=str(MANIFEST))
    parser.add_argument("--out", default=str(DEFAULT_OUTPUT))
    args = parser.parse_args()
    payload = json.loads(Path(args.manifest).read_text(encoding="utf-8"))
    build_document(payload, Path(args.out))


if __name__ == "__main__":
    main()
